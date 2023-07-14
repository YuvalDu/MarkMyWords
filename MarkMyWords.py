# --------------------------------------
# Filename:    MarkMyWords.py
# Built by:    RoniM, ShellyC, LiyaZ, YuvalD, HadarM
# Description: This script gets as an argument a directory,
#              and marks every docx file in it with a web beacon.
#              Causing the action of opening the file outside the closed domain,
#              in alerting our server of the action.
# ----------------------------------------

# Imports
import os
import re
import sys
import shutil
import hashlib
from docx import Document
from zipfile import ZipFile
from docx.shared import Inches

# Constants
SERVER_IP = "51.103.219.64"

# Change path to any .jpg/.png file
DUMMY_WATERMARK = "EyePic.jpg"


def create_watermark(file):
    """Create a tiny dummy watermark"""
    doc = Document(file)

    for section in doc.sections:
        section.header.is_linked_to_previous = False
        header = section.header
        header.paragraphs[0].alignment = 0
        run = header.paragraphs[0].add_run()
        run.add_picture(DUMMY_WATERMARK, width=Inches(0.0001), height=Inches(0.0001))

    doc.save(file)


def extract_docx_file(file):
    """Unzip docx file to a folder"""
    dest = (os.path.splitext(file))[0] + ".zip"
    os.rename(file, dest)
    os.mkdir(os.path.splitext(file)[0])
    with ZipFile(dest, 'r') as zObject:
        zObject.extractall(path=os.path.splitext(file)[0])
    os.remove(dest)


def convert_to_docx(dir):
    """Rezip all files to a docx file"""
    zip = shutil.make_archive(dir, 'zip', dir)
    dest = dir + ".docx"
    os.rename(zip, dest)
    shutil.rmtree(dir)


def update_xml(dir, hash):
    """Insert the url in a format of http://SERVERIP/FILENAME/HASH to every watermark related xml"""
    for xmlrel in os.listdir(fr"{dir}\word\_rels"):
        if xmlrel.startswith('header'):
            with open(fr"{dir}\word\_rels\{xmlrel}", 'r') as xml_file:
                data = xml_file.read()

            temp = data.split("Relationship ")
            filename = dir.split("\\")[1] + ".docx"
            url = fr'http://{SERVER_IP}/{filename}/{hash}'
            target = re.sub('Target=".*?"', fr'Target="{url}"', temp[1])
            external_target = temp[0] + 'Relationship TargetMode="External" ' + target

            with open(fr"{dir}\word\_rels\{xmlrel}", 'w') as xml_file:
                xml_file.write(external_target)


def main():
    target_dir = sys.argv[1]

    hash_dict = {}

    for file in os.listdir(target_dir):
        """
        For each file in given directory do:
        1. add name:hash to dict
        2. add dummy watermark
        3. extract him as a zip
        """
        hash_dict[file.split(".")[0]] = hashlib.md5(open(fr"{target_dir}\{file}", 'rb').read()).hexdigest()
        create_watermark(fr"{target_dir}\{file}")
        extract_docx_file(fr"{target_dir}\{file}")

    for file in os.listdir(target_dir):
        """Add a secret URL for each file, and rezip him to a docx file"""
        update_xml(fr"{target_dir}\{file}", hash_dict[file])
        convert_to_docx(fr"{target_dir}\{file}")


if __name__ == '__main__':
    main()
